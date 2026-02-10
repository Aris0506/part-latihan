from docx import Document
import re

def clean_text(text):
    # 1. Menghapus karakter aneh (non-alphanumeric) yang sering muncul di OCR
    # Tapi tetap menyisakan tanda baca dasar
    text = re.sub(r'[^\w\s.,?!\(\)\-\:\/]', '', text)
    
    # 2. Menghapus spasi ganda atau baris kosong berlebihan
    text = re.sub(r'\s+', ' ', text)
    
    # 3. Menghapus teks pendek yang biasanya noise (misal: "wm—x-vas")
    # Kita hapus kata yang isinya campuran huruf-angka aneh
    text = ' '.join([word for word in text.split() if len(word) > 1 or word.isalnum()])
    
    return text.strip()

def process_docx_cleaning(input_file):
    print(f"Memulai pembersihan: {input_file}")
    doc = Document(input_file)
    new_doc = Document()
    
    count = 0
    for para in doc.paragraphs:
        original_text = para.text
        if original_text.strip():
            cleaned = clean_text(original_text)
            # Hanya masukkan jika setelah dibersihkan teksnya masih bermakna
            if len(cleaned) > 5: 
                new_doc.add_paragraph(cleaned)
                count += 1
    
    output_file = input_file.replace(".docx", "_CLEANED.docx")
    new_doc.save(output_file)
    print(f"Selesai! {count} paragraf berhasil dibersihkan.")
    print(f"File bersih disimpan di: {output_file}")

# JALANKAN
if __name__ == "__main__":
    file_kamu = "MSIM4302 - Analisis Dan Perancangan Sistem_LENGKAP.docx" 
    process_docx_cleaning(file_kamu)