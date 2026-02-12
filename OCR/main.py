# import pytesseract
# from pdf2image import convert_from_path
# from docx import Document
# import os

# # 1. SETTING PATH (SESUAIKAN DENGAN KOMPUTERMU)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# PATH_POPPLER = r'C:\Users\USER\ais\Data Seminggu\OCR\poppler-25.12.0\Library\bin'

# def pdf_to_word(nama_pdf):
#     print(f"--- Memulai konversi: {nama_pdf} ---")
    
#     try:
#         # 2. Ubah PDF ke Gambar
#         print("Tahap 1: Mengubah PDF ke gambar...")
#         images = convert_from_path(nama_pdf, poppler_path=PATH_POPPLER)
        
#         doc = Document()
        
#         # 3. Proses tiap halaman
#         for i, image in enumerate(images):
#             print(f"Tahap 2: Membaca teks halaman {i+1}...")
#             # Pakai lang='ind' supaya akurat baca Bahasa Indonesia
#             text = pytesseract.image_to_string(image, lang='ind')
            
#             doc.add_heading(f'Halaman {i+1}', level=1)
#             doc.add_paragraph(text)
        
#         # 4. Simpan ke Word
#         nama_hasil = nama_pdf.replace(".pdf", ".docx")
#         doc.save(nama_hasil)
#         print(f"\nBERHASIL! Cek file: {nama_hasil}")

#     except Exception as e:
#         print(f"\nAda error nih bro: {e}")

# # JALANKAN
# if __name__ == "__main__":
#     # Taruh satu file PDF di folder yang sama, lalu ganti nama di bawah ini
#     file_target = "MSIM4302 - Analisis Dan Perancangan Sistem.pdf" 
    
#     if os.path.exists(file_target):
#         pdf_to_word(file_target)
#     else:
#         print(f"File '{file_target}' gak ketemu. Pastikan namanya benar dan satu folder ya!")



import pytesseract
from pdf2image import convert_from_path
from docx import Document
import os

# Konfigurasi Path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
PATH_POPPLER = r'C:\Users\USER\ais\Data Seminggu\OCR\poppler-25.12.0\Library\bin'

def ocr_bertahap(pdf_path, batch_size=10):
    print(f"--- Memulai OCR Bertahap: {pdf_path} ---")
    
    doc = Document()
    output_docx = pdf_path.replace(".pdf", "_LENGKAP.docx")
    
    # Kita proses sedikit-sedikit supaya RAM tidak penuh
    # Misal kita tes sampai halaman 476
    total_halaman = 493 
    
    for i in range(1, total_halaman + 1, batch_size):
        start = i
        end = min(i + batch_size - 1, total_halaman)
        
        print(f"\n>> Memproses Batch: Halaman {start} sampai {end}...")
        
        try:
            # Hanya ambil halaman tertentu (first_page & last_page)
            images = convert_from_path(
                pdf_path, 
                first_page=start, 
                last_page=end, 
                poppler_path=PATH_POPPLER
            )
            
            for j, image in enumerate(images):
                hal_skrg = start + j
                print(f"   OCR Halaman {hal_skrg}...")
                text = pytesseract.image_to_string(image, lang='ind')
                
                doc.add_heading(f'Halaman {hal_skrg}', level=1)
                doc.add_paragraph(text)
            
            # Simpan progress setiap batch selesai (biar kalau mati lampu, data gak hilang)
            doc.save(output_docx)
            
        except Exception as e:
            print(f"Error di halaman {start}-{end}: {e}")
            continue

    print(f"\nSelesai Bro! Cek file: {output_docx}")

# Jalankan
if __name__ == "__main__":
    file_pdf = "Basis Data.pdf" # Sesuaikan nama filenya
    ocr_bertahap(file_pdf)





# import pytesseract
# from pdf2image import convert_from_path
# from docx import Document
# import os

# # Konfigurasi Path
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# PATH_POPPLER = r'C:\Users\USER\ais\Data Seminggu\OCR\poppler-25.12.0\Library\bin'

# def ocr_tambal(pdf_path, batch_size=5):
#     print(f"--- Memulai Tambal OCR: {pdf_path} ---")
    
#     doc = Document()
#     # Beri nama file baru agar tidak bentrok (Permission Denied)
#     output_docx = "TAMBALAN_Halaman_601_614.docx"
    
#     # KUNCI PERUBAHAN: Mulai dari 601 sampai 614
#     start_hal = 601
#     total_halaman = 614 
    
#     for i in range(start_hal, total_halaman + 1, batch_size):
#         start = i
#         end = min(i + batch_size - 1, total_halaman)
        
#         print(f"\n>> Memproses Batch: Halaman {start} sampai {end}...")
        
#         try:
#             images = convert_from_path(
#                 pdf_path, 
#                 first_page=start, 
#                 last_page=end, 
#                 poppler_path=PATH_POPPLER
#             )
            
#             for j, image in enumerate(images):
#                 hal_skrg = start + j
#                 print(f"   OCR Halaman {hal_skrg}...")
#                 text = pytesseract.image_to_string(image, lang='ind')
                
#                 doc.add_heading(f'Halaman {hal_skrg}', level=1)
#                 doc.add_paragraph(text)
            
#             doc.save(output_docx)
            
#         except Exception as e:
#             print(f"Error di halaman {start}-{end}: {e}")
#             continue

#     print(f"\nSelesai Bro! Hasil tambalan ada di: {output_docx}")

# # Jalankan
# if __name__ == "__main__":
#     file_pdf = "Algoritma Dan Pemrograman.pdf" 
#     # Pastikan file Word utama sudah ditutup sebelum menjalankan ini
#     ocr_tambal(file_pdf)